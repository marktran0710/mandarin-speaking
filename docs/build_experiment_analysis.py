from pathlib import Path
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
REFERENCE = Path(r"C:\Users\Administrator\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.1\skills\artifact-template-experiment-analysis\assets\reference.docx")
OUTPUT = ROOT / "output" / "experiment-analysis-irt-practice.docx"

paragraph_replacements = {
    "Experiment ": "Mandarin Speaking Practice",
    "Report Name": "Measurement, IRT Question Bank, and Practice UI Experiment Analysis",
    "Author": "Research / Product Team",
    "Month DD, YYYY": "August 2026",
    "[Describe what this document records, including the experiment objective, execution, analysis, and decision it supports. Note the product surface and intended scope.]": "This document defines the measurement and analysis plan for the teacher dashboard, IRT question bank, and revised student pronunciation-practice experience. It is a pre-analysis document; no treatment effect or rollout decision is claimed yet.",
    "[Explain why this record is being maintained, who should use it, and how it supports auditability, replication, or future decision-making.]": "The research, product, and teaching teams use this record to preserve metric definitions, event provenance, IRT assumptions, and analysis decisions before data collection.",
    "[Describe the company, product, user journey, and observed problem or opportunity. Summarize the evidence that motivated the experiment and where the issue appears in the experience.]": "Students currently receive pronunciation scores, but may not know whether a result is a real tone error, insufficient audio evidence, or an ASR mismatch. Teachers need class-level visibility into these states and repeated error patterns.",
    "[Explain why this test surface and intervention were selected, including expected value, operational risk, reversibility, and any important alternatives that remained unchanged.]": "The intervention targets feedback clarity and adaptive item selection because both are reversible and directly connected to measurable learning behavior. Audio analysis, backend tone scoring, authentication, and teacher-authored content remain unchanged unless explicitly instrumented.",
    "[Describe the experimental design and the single change being tested. Explain how the design supports causal interpretation and which elements remained unchanged.]": "The future comparison is between the current practice flow and a revised flow that makes evidence status explicit, highlights the next actionable word, and selects calibrated questions near the learner's estimated ability. The current implementation phase only instruments and validates the surfaces; it does not claim a treatment effect.",
    "[Describe eligibility, assignment, exclusions, runtime, and other choices made to reduce contamination or bias. Note any platform, geography, or device constraints.]": "Eligibility, class-level assignment, exclusion rules, and runtime are intentionally left as pre-registration decisions for the later pilot. Invalid audio and unjudged attempts will be excluded from pronunciation outcome denominators but retained as system-quality observations.",
    "[Write an if/then hypothesis naming the control, treatment, audience, expected metric direction, and causal rationale.]": "If students receive evidence-labeled, next-action pronunciation feedback and questions matched to their estimated ability, then judged pronunciation pass rate and delayed retention should improve without increasing not-enough-evidence rates or abandonment.",
    "[Define the primary success metric, guardrail metrics, decision thresholds, and the conditions required to declare a winner or proceed.]": "Primary metric: judged pronunciation mastery pass rate. Guardrails: not-enough-evidence rate, recording retry rate, time to pass, completion rate, and ASR/audio failure rate. Thresholds and statistical decision rules will be locked before the pilot.",
    "[Confirm that only the intended treatment differs between variants.]": "Validation item: compare rendered UI, question exposure, reference audio, scoring version, and backend response fields between conditions.",
    "[Confirm assignment persistence, randomization unit, and allocation logic.]": "Validation item: persist condition at the class or student assignment unit and verify that it does not change across sessions.",
    "[Validate exposure, outcome, and guardrail instrumentation.]": "Validation item: verify event names, schema version, attempt identifiers, item identifiers, and outcome denominators against raw recordings and quiz responses.",
    "[Document traffic exclusions and data-quality filters.]": "Exclude duplicate events, test/admin sessions, missing item IDs from IRT calibration, and pronunciation outcomes with no measurable syllables. Keep excluded rows in an audit table.",
    "[Record staging, A/A, or other pre-launch validation results.]": "Pending implementation validation. The current acceptance gate is schema validation plus focused UI and utility tests.",
    "[Summarize achieved sample size, exposure balance, runtime, and any sample-ratio or eligibility findings. State whether the data are suitable for analysis.]": "No experimental exposure has been run in this implementation phase. Data are suitable for instrumentation and calibration checks only, not for a causal conclusion.",
    "[Summarize the primary metric result, uncertainty, statistical significance, and practical effect size. Avoid claims beyond the experiment scope.]": "Not estimable before the pilot. Report the numerator, denominator, cluster structure, confidence interval, and effect size after data collection.",
    "[Summarize each guardrail result, note any threshold breaches, and explain whether the risk profile supports the proposed decision.]": "Not estimable before the pilot. The dashboard will separate learning failures from audio-quality failures so guardrails are not conflated.",
    "[Describe which segments were pre-specified, why they matter, and whether the experiment was powered for subgroup comparisons. Distinguish confirmatory from directional analysis.]": "Pre-specified descriptive segments: proficiency level, class, tone pattern, question difficulty, device type, and reference-audio source. Unless powered separately, subgroup results remain directional.",
    "[Summarize the most decision-relevant segment patterns and plausible mechanisms. Note uncertainty, multiplicity, or sample-size caveats.]": "Pending data. Avoid interpreting small subgroup differences as treatment effects.",
    "[List concurrent launches, operational events, or external factors that could affect interpretation.]": "Track model changes, backend scoring changes, teacher material edits, school-calendar effects, and microphone/device changes.",
    "[Describe traffic-mix or eligibility changes during the experiment.]": "Record changes to class enrollment, proficiency mix, attendance, and assignment exposure.",
    "[Document logging, pipeline, or warehouse incidents and their resolution.]": "Record missing events, delayed analysis responses, duplicate attempt IDs, and any backfill.",
    "[State platform, geography, audience, or time-window limitations.]": "Initial evidence will be limited to the participating classes, supported browser/device combinations, Taiwan Mandarin targets, and the observed study window.",
    "[Note treatment-specific risks, policy constraints, or generalizability limits.]": "Do not treat unjudged audio as a student failure. Do not generalize IRT parameters across levels or dialect targets without recalibration.",
    "[Explain how these limitations affect the recommendation and clearly define where the evidence should and should not be generalized.]": "Recommendations apply first to the instrumented practice flow and participating population. Broader rollout requires replication and monitoring.",
    "[Connect the observed results to the hypothesis and proposed mechanism. Explain whether the effect is statistically credible, practically meaningful, and consistent across key measures.]": "Pending pilot results. Interpretation must combine the primary outcome, guardrails, confidence intervals, class clustering, and data-quality exclusions.",
    "[Translate the result into business or product impact, including expected upside, cost, reversibility, and remaining uncertainty. State the recommended decision posture.]": "Pending pilot results. The default posture is staged validation because the UI and IRT changes are reversible but incorrect feedback can harm learner trust.",
    "[Describe the rollout or implementation sequence.]": "After pilot analysis: fix instrumentation issues, recalibrate items, review teacher dashboard usefulness, then consider a controlled production rollout.",
    "[Define post-launch monitoring metrics, cadence, and duration.]": "Monitor weekly: judged pass rate, not-enough-evidence rate, retry rate, item exposure, item difficulty drift, event completeness, and teacher intervention time.",
    "[List replication or follow-up experiments.]": "Follow-ups may test reference-audio quality, explicit teacher phrase boundaries, and adaptive difficulty separately.",
    "[Identify additional analysis or product opportunities.]": "Evaluate delayed retention, transfer to unseen sentences, DIF by proficiency/device, and teacher actionability of dashboard recommendations.",
    "[Document archival, reporting, and ownership requirements.]": "Archive raw events, calibration snapshots, analysis code version, rendered dashboard evidence, and the final reviewed report.",
    "[Metric name 1]: [Define numerator, denominator, eligibility, event source, and calculation.]": "Judged mastery pass rate: passed pronunciation mastery / attempts with at least one judged syllable; source: analysis_completed event and backend pronunciation_mastery.",
    "[Metric name 2]: [Define numerator, denominator, eligibility, event source, and calculation.]": "Not-enough-evidence rate: attempts marked retry or not_judged / analyzed attempts; source: feedback quality and pronunciation mastery.",
    "[Metric name 3]: [Define numerator, denominator, eligibility, and calculation.]": "Time to pass: elapsed time from practice_started to first practice_passed event for the same student, scene, and target.",
    "[Metric name 4]: [Define numerator, denominator, eligibility, and calculation.]": "Item difficulty: logit transform of clipped proportion correct during calibration; exclude items with fewer than 30 usable responses from publishing.",
    "[Metric name 5]: [Define numerator, denominator, eligibility, and calculation.]": "Event completeness: expected instrumented transitions observed / expected transitions from submitted attempts; report separately from learning outcomes.",
    "[Document the statistical method, confidence level, and software or query version used.]": "Planned analysis: cluster-aware comparison at the class level, with confidence intervals and effect sizes; IRT calibration uses the versioned utility contract and item-response audit table.",
    "[Record stopping rules, interim reads, or deviations from the analysis plan.]": "No interim stopping rule is active in the implementation phase. Any later deviation must be timestamped and documented before reading treatment outcomes.",
    "[Summarize data reconciliation, backfills, and reviewer sign-off.]": "Reconcile dashboard totals with raw event storage, audio records, and quiz attempts. Obtain research and teacher-review sign-off before using the report for a rollout decision.",
    "[Add any remaining assumptions, caveats, or provenance notes.]": "This document contains no invented results. All displayed outcomes must be populated from observed, versioned data.",
}

table_replacements = {
    "[Version]": "0.1",
    "[Prepared by / team]": "Research / Product Team",
    "[Reviewer roles or names]": "Pending",
    "[Date prepared]": "2026-08-11",
    "[Reporting start date] to [Reporting end date]": "Implementation phase",
    "[Draft / In review / Final]": "Draft",
    "[Experiment name]": "Evidence-labeled pronunciation practice",
    "[Experiment key]": "mandarin-practice-evidence-irt-v1",
    "[Owner team]": "Learning Product",
    "[Business owner]": "Mandarin Speaking Team",
    "[Product surface]": "Student practice + teacher dashboard",
    "[Primary objective]": "Improve actionable feedback and adaptive question selection",
    "[Describe the control experience]": "Current student practice flow and fixed question selection",
    "[Describe the treatment experience]": "Evidence-labeled feedback, next-action guidance, and calibrated item selection",
    "[Variant allocation]": "To be pre-registered later; not run in this phase",
    "[Randomization unit]": "Class or student, to be locked before pilot",
    "[Eligible audience]": "Mandarin learners using the supported student practice flow",
    "[Traffic or user exclusions]": "Admin/test sessions, invalid audio, duplicate events",
    "[Start date]": "TBD",
    "[End date]": "TBD",
    "[Planned runtime]": "TBD",
    "[Actual runtime]": "Not started",
    "[Define the primary metric threshold and statistical decision rule]": "Judged pronunciation mastery pass rate; lock threshold before pilot",
    "[Define guardrail metric and acceptable threshold]": "Not-enough-evidence rate; must not worsen materially",
    "[Define guardrail metric and acceptable threshold]": "Retry/completion rate; monitor for harm",
    "[State the full launch, hold, or rollback rule]": "Hold until primary metric and guardrails pass pre-registered rules",
}

def replace_text(text: str) -> str:
    if text in paragraph_replacements:
        return paragraph_replacements[text]
    if text in table_replacements:
        return table_replacements[text]
    return text

def set_cell(cell, text):
    if cell.text == text:
        cell.text = replace_text(text)

doc = Document(REFERENCE)
for paragraph in doc.paragraphs:
    if paragraph.text in paragraph_replacements:
        paragraph.text = paragraph_replacements[paragraph.text]
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            set_cell(cell, cell.text)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
