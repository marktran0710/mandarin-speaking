from docx import Document
from pathlib import Path
import shutil

ref = Path(r"C:\Users\Administrator\.codex\plugins\cache\openai-curated-remote\openai-templates\0.1.1\skills\artifact-template-experiment-analysis\assets\reference.docx")
out = Path(r"D:\hautran\Lab\mandarin-speaking\IRT_Experimental_Design_Approaches.docx")
shutil.copy2(ref, out)
doc = Document(out)

def setp(i, text):
    p = doc.paragraphs[i]
    p.text = text

setp(0, "Experiment Report")
setp(6, "IRT Experimental Design")
setp(7, "Control and Grouping Approaches for a Mandarin-Speaking System")
setp(23, "Author: Product and Learning Analytics Team")
setp(24, "August 13, 2026")

texts = {
29: "Purpose",
30: "This document presents three real-world ways to collect item-response data and control experiments in a Mandarin-speaking learning system. It covers how to define control and treatment groups, protect IRT calibration, and decide whether an intervention should scale.",
31: "Use this as an experiment-design playbook before launch. Each approach separates the learning intervention from measurement, preserves a stable comparison group, and records enough provenance to reproduce the analysis.",
33: "Business Context",
34: "The system delivers Mandarin speaking prompts, evaluates responses, and may adapt difficulty, feedback, or practice sequencing. The core measurement need is to estimate learner ability and item difficulty from reliable item-level responses while testing product changes.",
35: "The main risk is allowing a new coaching or scoring experience to change response behavior in a way that looks like ability growth. The safest design keeps the IRT item bank, scoring rules, and exposure logging stable unless those are the explicit treatment variables.",
36: "Experiment Summary",
39: "Design Notes",
40: "Use one primary treatment contrast at a time. Control should receive the current production experience; treatment receives the new feedback, sequencing, or interface. Log learner_id, item_id, prompt version, response, score, latency, attempt number, assignment, exposure timestamp, and model version.",
41: "Randomize at the level that prevents contamination: learner for independent experiences, learner-by-period for crossover tests, or classroom/teacher/site for shared environments. Keep anchor items in both groups so IRT parameters remain linked across conditions.",
43: "Hypothesis",
44: "If learners receive the treatment experience, then their post-test ability estimate or targeted speaking outcome will improve versus control, without unacceptable degradation in response validity, completion, latency, or item exposure balance.",
46: "Success Criteria",
47: "Primary: pre-registered improvement in post-test ability or task success, with a confidence interval and minimum practically important effect. Guardrails: missing-response rate, scoring disagreement, latency, dropout, item exposure concentration, and sample-ratio mismatch. Ship only if the primary rule passes and all critical guardrails pass.",
49: "Pre-Launch Validation",
50: "Run an A/A test to confirm the assignment pipeline produces equivalent outcomes.",
51: "Verify sticky assignment, the randomization unit, allocation, and no cross-group leakage.",
52: "Validate exposure, response, audio, transcript, scoring, and IRT-event instrumentation end to end.",
53: "Exclude bots, duplicate sessions, test accounts, corrupted audio, and learners without consent where applicable.",
54: "Freeze the analysis plan, anchor-item list, stopping rule, and parameter version before reading outcomes.",
56: "Sample and Exposure Summary",
57: "Report learners, item responses, unique items, anchor-item coverage, allocation share, runtime, and sample-ratio checks. For IRT, report whether each item has enough responses across ability levels and whether calibration and validation data were kept distinct.",
59: "Primary Outcome",
61: "Use one primary outcome per experiment: change in held-out speaking-task success, post-test ability estimate, or a pre-specified latent-trait difference. Report effect size, uncertainty, practical threshold, and whether the result is identified by independent validation items.",
63: "Guardrail Metrics",
65: "Monitor scoring reliability, audio/transcript failure, response latency, learner dropout, item exposure, anchor-item drift, and differential item functioning signals. A treatment that raises scores by changing scoring behavior rather than learner performance should not be declared a win.",
66: "Segment Results",
67: "Pre-specify segments such as baseline ability band, course level, device type, learning history, and first-language background when ethically and operationally appropriate. Treat subgroup results as confirmatory only when powered for them.",
69: "Inspect whether the effect is consistent across ability bands and whether any item shows suspicious treatment-specific behavior. Use directional segment findings to design the next experiment, not to overrule the primary decision without sufficient power.",
71: "Data Quality and Limitations",
72: "Record concurrent curriculum changes, speech-model releases, outages, teacher interventions, and major traffic events.",
73: "Track changes in learner mix, eligibility, device mix, and practice intensity over time.",
74: "Reconcile event counts with warehouse tables and document backfills, scoring changes, or missing audio.",
75: "Results may not generalize beyond the tested language level, task type, geography, device, or runtime.",
76: "IRT estimates can be biased by item leakage, unbalanced exposure, local dependence, speeded responses, or differential item functioning.",
77: "These limits affect how broadly the treatment can be rolled out. Preserve the raw response data and item versions so later calibration or replication can separate product effects from measurement effects.",
79: "Interpretation",
80: "Interpret the outcome jointly with IRT diagnostics. A credible result should improve the target outcome on held-out or anchor-linked items, remain plausible across ability bands, and avoid material shifts in discrimination, guessing, or item characteristic behavior unless those shifts were intended.",
81: "Translate the effect into learner and business impact, including expected learning gain, implementation cost, reversibility, and uncertainty. Prefer staged rollout with continued anchor monitoring when the treatment changes feedback or item selection.",
83: "Decision Log",
86: "Post Test Actions",
87: "Roll out gradually: 10% exposure, then 25%, 50%, and 100% only after guardrails remain stable.",
88: "Monitor primary outcome, scoring reliability, anchor-item drift, exposure concentration, and dropout weekly for four to eight weeks.",
89: "Replicate the result on a fresh learner cohort and a fresh set of held-out items.",
90: "Follow up on suspicious items, subgroup heterogeneity, and any evidence of practice or scoring contamination.",
91: "Archive the assignment seed, analysis query, item-bank version, model version, exclusions, and reviewer sign-off.",
93: "Appendix A: Metric Definitions",
94: "Primary outcome: pre-specified held-out task success or change in IRT ability, with numerator, denominator, item set, and estimation version documented.",
95: "Response validity: valid audio and transcript responses divided by exposed speaking tasks.",
96: "Anchor drift: change in anchor-item difficulty or characteristic-curve behavior relative to the frozen reference calibration.",
97: "Exposure concentration: share of all responses contributed by the most-exposed items, compared with the planned cap.",
98: "Scoring reliability: agreement or correlation between automated scoring and human/reference scoring on the validation sample.",
100: "Appendix B: Analyst Notes",
101: "Use a pre-registered IRT model and estimation method; state confidence level, software/query version, and whether parameters are fixed, recalibrated, or jointly estimated.",
102: "Do not stop on a promising interim result unless the stopping rule allows it. Record every interim read and any deviation from the analysis plan.",
103: "Reconcile assignment, exposure, response, scoring, and warehouse counts before analysis; obtain independent review of the final query and decision.",
104: "Keep raw responses immutable and version all item text, audio prompts, scoring models, and feature flags.",
}
for i,t in texts.items(): setp(i,t)

# Tables: concise operational content
tables = doc.tables
def fill(tbl_i, rows):
    t = tables[tbl_i]
    for r,row in enumerate(rows):
        if r >= len(t.rows): t.add_row()
        for c,val in enumerate(row):
            if c < len(t.rows[r].cells): t.rows[r].cells[c].text = val

fill(0, [["Version","1.0"],["Prepared By","Product and Learning Analytics Team"],["Reviewers","Learning science, data science, engineering"],["Date Prepared","August 13, 2026"],["Reporting Window","Design guidance; apply per experiment"],["Status","Draft playbook"]])
fill(1, [["Field","Details"],["Experiment Name","IRT design-control playbook"],["Experiment Key","IRT-DESIGN-001"],["Owner Team","Learning Analytics"],["Business Owner","Mandarin Learning Product"],["Product Surface","Speaking practice and feedback"],["Primary Objective","Estimate learning impact without corrupting measurement"],["Control (Variant A)","Current production experience"],["Treatment (Variant B)","One new feedback, sequencing, or interface change"],["Allocation","50/50 default; stage rollout after decision"],["Unit of Randomization","Learner, learner-by-period, or cluster"],["Audience","Eligible consenting learners"],["Exclusions","Bots, duplicates, test accounts, corrupted sessions"],["Start Date","Set per experiment"],["End Date","Set per experiment"],["Planned Runtime","Until power and exposure criteria are met"],["Actual Runtime","Record after close"]])
fill(2, [["Metric","Target/Rule"],["Primary Metric","Pre-registered held-out outcome or post-test ability; CI and practical threshold"],["Guardrail 1","Scoring reliability and validity: no material breach"],["Guardrail 2","Dropout, latency, and missing audio: within agreed limits"],["Guardrail 3","Anchor drift and item exposure: within caps"],["Decision Rule","Launch only when primary passes and critical guardrails pass; otherwise hold, iterate, or roll back"]])
fill(3, [["Metric","Variant A (Control)","Variant B (Treatment)"],["Calibration responses","Production calibration stream","Separate calibration/validation stream"],["Held-out outcome","Measure on frozen items","Measure on frozen items"],["Anchor coverage","Same anchor set","Same anchor set"],["Assignment share","50% target","50% target"],["Sample Ratio Check","Pass when observed allocation matches plan","Pass when observed allocation matches plan"]])
fill(4, [["Measure","Variant A","Variant B","Absolute Delta","Relative Delta"],["Held-out learning outcome","Estimate","Estimate","B-A","(B-A)/A"],["95% Confidence Interval","N/A","N/A","Report interval","N/A"],["Two-Sided p-value","N/A","N/A","Report p-value","N/A"],["Outcome Narrative","Current experience","New experience","Interpret practical effect","Avoid causal claims outside scope"]])
fill(5, [["Guardrail Metric","Variant A","Variant B","Delta","Status"],["Scoring reliability","Report","Report","B-A","Pass/Fail"],["Missing audio/transcript","Report","Report","B-A","Pass/Fail"],["Dropout/latency","Report","Report","B-A","Pass/Fail"],["Anchor drift/exposure","Report","Report","B-A","Pass/Fail"]])
fill(6, [["Segment","Control Value","Treatment Value","Absolute Delta","Interpretation"],["Baseline ability band","Report","Report","B-A","Pre-specified only"],["Course level","Report","Report","B-A","Pre-specified only"],["Device type","Report","Report","B-A","Directional unless powered"],["New vs returning learner","Report","Report","B-A","Check contamination"],["Task/item family","Report","Report","B-A","Inspect item behavior"]])
fill(7, [["Item","Decision"],["Final Decision","Choose rollout, hold, iterate, or rollback based on primary and guardrails"],["Decision Date","Record date"],["Approvers","Product, learning science, data science"],["Rollout Type","Staged exposure with continued anchor monitoring"],["Rollback Trigger","Critical guardrail breach, validity failure, or unexplained anchor drift"],["Follow Up","Fresh cohort replication and held-out item validation"]])

doc.save(out)
print(out)
